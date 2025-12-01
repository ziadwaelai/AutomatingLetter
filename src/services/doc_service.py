from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from docx.oxml.ns import qn
import os
import json
from typing import Dict, Optional, Union
import logging
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from dotenv import load_dotenv
import requests
from PIL import Image

# Load environment variables from .env file
load_dotenv()

logger = logging.getLogger(__name__)


class CreateDocument:
    """
    A class to create Word documents from JSON input.
    Takes JSON with date, title, body, besmallah, and footer fields
    and generates a formatted Word document based on the Netzero template.
    """

    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize the CreateDocument service.

        Args:
            template_path: Path to the Netzero.docx template.
                          If None, tries environment variable or relative path.
        """
        if template_path is None:
            # Try to get template path from environment variable
            env_template_path = os.getenv("DOCX_TEMPLATE_PATH")

            if env_template_path:
                self.template_path = env_template_path
                logger.info(f"Using template path from DOCX_TEMPLATE_PATH env var: {self.template_path}")
            else:
                # Use relative path from current directory (server-safe)
                # This will work on any server regardless of location
                base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                self.template_path = os.path.join(base_dir, "LetterToPdf", "templates", "Netzero.docx")
                logger.info(f"Using default template path: {self.template_path}")

            # If template doesn't exist, log warning and set to None (will create blank document)
            if not os.path.exists(self.template_path):
                logger.warning(f"Template file not found at: {self.template_path}, will create blank document")
                self.template_path = None
        else:
            self.template_path = template_path
            # If custom path doesn't exist, log warning but allow blank document creation
            if not os.path.exists(self.template_path):
                logger.warning(f"Custom template path not found: {self.template_path}, will create blank document")
                self.template_path = None

        self.document = None
        self.letter_id = None  # Store letter ID for the document

    def load_template(self) -> Document:
        """
        Load the template document.

        Returns:
            Loaded Document object
        """
        try:
            self.document = Document(self.template_path)
            logger.info(f"Template loaded successfully from: {self.template_path}")
            return self.document
        except Exception as e:
            logger.error(f"Failed to load template: {e}")
            raise

    def _parse_raw_letter_content(self, letter_content: str) -> Dict[str, str]:
        """
        Parse raw letter content using OpenAI GPT-4 to extract structured fields.

        Args:
            letter_content: Raw letter text containing all content in one string

        Returns:
            Dictionary with extracted fields: title, body, besmallah, footer, date

        Raises:
            ValueError: If parsing fails or OpenAI API is not available
        """
        try:
            # Try to import config, with fallback for direct script execution
            try:
                from ..config import get_config
                config = get_config()
            except (ImportError, ValueError):
                # Fallback: try to get from environment for direct execution
                import os
                openai_api_key = os.getenv("OPENAI_API_KEY")
                if not openai_api_key:
                    raise ValueError(
                        "OpenAI API key not found. Please set OPENAI_API_KEY environment variable "
                        "or configure it in your settings."
                    )
                # Create a simple config object
                class SimpleConfig:
                    pass
                config = SimpleConfig()
                config.openai_api_key = openai_api_key

            if not config.openai_api_key:
                raise ValueError("OpenAI API key is not configured")

            # Initialize LLM with GPT-4o (latest model with better parsing)
            llm = ChatOpenAI(
                model="gpt-4o",
                temperature=0.3,
                openai_api_key=config.openai_api_key,
                timeout=30,
            )

            # Create parsing prompt
            prompt = PromptTemplate.from_template("""
أنت مساعد متخصص في تحليل وتنسيق الخطابات الرسمية العربية.

قم بتحليل الخطاب التالي واستخرج المعلومات التالية:

**الخطاب المراد تحليله:**
{letter_content}

استخرج المعلومات التالية وأعد الاستجابة بتنسيق JSON صحيح:

1. **besmallah**: الفقرة الأولى أو بداية الخطاب (عادة "بسم الله الرحمن الرحيم" أو ما يشابهها). إذا لم توجد، استخدم سطر فارغ ""
2. **title**: عنوان الخطاب الرئيسي (REQUIRED - إلزامي). ابحث عن العنوان في:
   - السطور الأولى بعد البسملة
   - أي نص مكتوب بخط عريض أو مميز
   - الموضوع الرئيسي للخطاب
   - إذا لم تجد عنوان واضح، استنتج عنوانًا مناسبًا من محتوى الخطاب (مثل: "طلب...", "إفادة...", "خطاب رسمي...")
   - **يجب أن يكون العنوان موجودًا دائمًا ولا يمكن أن يكون فارغًا**
3. **date**: التاريخ إن وجد (قد يكون في البداية أو النهاية). إذا لم يوجد، استخدم سطر فارغ ""
4. **body**: محتوى الخطاب الرئيسي (REQUIRED - إلزامي). يشمل كل المحتوى بين العنوان والختام
5. **footer**: الختام والتوقيع (عادة في نهاية الخطاب مثل "وتفضلوا بقبول فائق الاحترام والتقدير"). إذا لم يوجد، استخدم سطر فارغ ""

⚠️ **مهم جداً:**
- الحقول **title** و **body** إلزامية ويجب أن تحتوي على قيمة
- إذا لم تجد عنوانًا صريحًا، استنتج عنوانًا من المحتوى
- الحفاظ على الترتيب الطبيعي للخطاب
- استخراج كل جزء بشكل كامل ودقيق
- عدم ترك أي محتوى مهم خارج body
- الحفاظ على التنسيق والأسطر الجديدة

أعد الاستجابة بصيغة JSON فقط (بدون أي نص إضافي):
{{
    "besmallah": "...",
    "title": "...",
    "date": "...",
    "body": "...",
    "footer": "..."
}}
""")

            # Create parser
            parser = JsonOutputParser()

            # Create chain
            chain = prompt | llm | parser

            logger.info("Starting raw letter content parsing with GPT-4o")

            # Parse the content
            result = chain.invoke({"letter_content": letter_content})

            logger.info("Letter content parsed successfully")
            logger.debug(f"Parsed fields: {list(result.keys())}")

            # Validate that required fields are present and not empty
            if not result.get('title') or not result.get('title').strip():
                # Extract a title from the first non-empty line of body as fallback
                body_lines = result.get('body', '').split('\n')
                extracted_title = None
                for line in body_lines:
                    line = line.strip()
                    if line and len(line) < 100:  # Use first short line as title
                        extracted_title = line
                        break

                if extracted_title:
                    result['title'] = extracted_title
                    logger.warning(f"Title was empty, extracted from body: {extracted_title}")
                else:
                    result['title'] = "خطاب رسمي"
                    logger.warning("Title was empty, using default: خطاب رسمي")

            if not result.get('body') or not result.get('body').strip():
                logger.error("Body field is empty after parsing")
                raise ValueError("Failed to extract body content from letter")

            return result

        except Exception as e:
            logger.error(f"Failed to parse letter content: {e}")
            raise ValueError(f"Failed to parse letter content using OpenAI: {str(e)}")

    def create_from_raw_content(self, letter_content: str) -> Document:
        """
        Create a document from raw letter content.
        This method uses OpenAI to parse the raw content and extract structured fields.

        Args:
            letter_content: Raw letter text as a single string

        Returns:
            Document object

        Raises:
            ValueError: If parsing fails or required fields are missing
        """
        logger.info("Creating document from raw letter content")

        # Parse the raw content using GPT-4
        parsed_data = self._parse_raw_letter_content(letter_content)

        # Validate that we have at least title and body
        if not parsed_data.get("title") or not parsed_data.get("body"):
            raise ValueError("Failed to extract title or body from the letter content")

        # Use the structured create_from_json method
        return self.create_from_json(parsed_data)

    def create_from_json(self, json_input: Union[str, Dict], letter_id: Optional[str] = None) -> Document:
        """
        Create a document from JSON input.

        Args:
            json_input: Either a JSON string or a dictionary containing:
                       - besmallah: Opening Islamic greeting (optional)
                       - title: Document title
                       - date: Document date
                       - body: Main content of the document
                       - footer: Footer text (optional)
                       - include_signature: Boolean flag to include signature section (optional)
                       - signature_image_url: Google Drive URL for signature image (required if include_signature=True)
                       - signature_name: Name of the signer (required if include_signature=True)
                       - signature_job_title: Job title of the signer (required if include_signature=True)
                       - signature_section_title: Custom title for signature section (optional, default: "التوقيع")
            letter_id: Optional letter ID to embed in the document (LET-YYYYMMDD-XXXXX format)

        Returns:
            Document object
        """
        # Parse JSON if string is provided
        if isinstance(json_input, str):
            try:
                data = json.loads(json_input)
            except json.JSONDecodeError as e:
                logger.error(f"Invalid JSON input: {e}")
                raise ValueError(f"Invalid JSON input: {e}")
        else:
            data = json_input

        # Check if letter_id is in the data dictionary
        if letter_id is None and isinstance(data, dict) and 'letter_id' in data:
            letter_id = data.get('letter_id')

        # Store letter_id for later use
        if letter_id:
            self.letter_id = letter_id
            logger.info(f"Document letter ID set to: {letter_id}")

        # Validate required fields
        required_fields = ['title', 'body']
        missing_fields = [field for field in required_fields if field not in data or not data[field]]
        if missing_fields:
            raise ValueError(f"Missing required fields: {', '.join(missing_fields)}")

        # Load template
        self.load_template()

        # Clear existing content (keep structure)
        self._clear_document_content()

        # Add document info header (date and letter ID) if letter_id is provided
        if self.letter_id:
            self.add_document_info_header(self.letter_id)

        # Add besmallah if provided
        if data.get('besmallah'):
            self.add_besmallah(data['besmallah'])

        # Add title
        self.add_title(data['title'])

        # Add body content
        self.add_body(data['body'])

        # Add footer if provided
        if data.get('footer'):
            self.add_footer(data['footer'])

        # Add signature section if requested
        if data.get('include_signature', False):
            signature_image_url = data.get('signature_image_url', '').strip()
            signature_name = data.get('signature_name', '').strip()
            signature_job_title = data.get('signature_job_title', '').strip()
            signature_section_title = data.get('signature_section_title', 'التوقيع').strip()

            # Validate signature fields
            if not signature_image_url or not signature_name or not signature_job_title:
                logger.warning(
                    "include_signature is True but missing required signature fields "
                    "(signature_image_url, signature_name, signature_job_title). Skipping signature section."
                )
            else:
                # Add signature section
                self.add_signature_section(
                    signature_image_url=signature_image_url,
                    job_title=signature_job_title,
                    name=signature_name,
                    section_title=signature_section_title
                )

        logger.info(f"Document created successfully from JSON (Letter ID: {self.letter_id})")
        return self.document

    def _clear_document_content(self):
        """
        Clear the document content while preserving structure.
        Removes all paragraphs and tables after the first one.
        """
        try:
            # Keep the first paragraph (style preservation)
            # Remove all other paragraphs
            for paragraph in self.document.paragraphs[1:]:
                p = paragraph._element
                p.getparent().remove(p)

            # Remove all tables
            for table in self.document.tables:
                tbl = table._element
                tbl.getparent().remove(tbl)

            logger.debug("Document content cleared")
        except Exception as e:
            logger.warning(f"Error clearing document content: {e}")

    def add_document_info_header(self, letter_id: str) -> None:
        """
        Add document info header (date and letter ID) matching PDF style.
        Displays Gregorian date, Hijri date, and letter ID.

        Args:
            letter_id: The letter ID (e.g., "LET-20251105-12345")
        """
        # Get current dates in both formats
        from hijri_converter import Gregorian
        import pytz
        from datetime import datetime

        try:
            # Get current date in KSA timezone
            ksa_tz = pytz.timezone('Asia/Riyadh')
            now_ksa = datetime.now(ksa_tz)

            # Gregorian date in Arabic
            gregorian_arabic = f"{now_ksa.day} / {now_ksa.month} / {now_ksa.year}"

            # Hijri date
            hijri_date = Gregorian(now_ksa.year, now_ksa.month, now_ksa.day).to_hijri()
            hijri_arabic = f"{hijri_date.day} / {hijri_date.month} / {hijri_date.year}"
        except Exception as e:
            logger.warning(f"Error getting dates: {e}")
            gregorian_arabic = ""
            hijri_arabic = ""

        # Add Gregorian date
        para1 = self.document.add_paragraph()
        para1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run1 = para1.add_run(f"التاريخ: {gregorian_arabic} ميلادي")
        run1.font.size = Pt(8)
        run1.font.name = 'Tajawal'
        para1.paragraph_format.space_before = Pt(0)
        para1.paragraph_format.space_after = Pt(0)
        para1.paragraph_format.line_spacing = 1.0

        # Add Hijri date
        para2 = self.document.add_paragraph()
        para2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run2 = para2.add_run(f"الموافق: {hijri_arabic} هجري")
        run2.font.size = Pt(8)
        run2.font.name = 'Tajawal'
        para2.paragraph_format.space_before = Pt(0)
        para2.paragraph_format.space_after = Pt(0)
        para2.paragraph_format.line_spacing = 1.0

        # Add Letter ID - format with proper RTL order
        para3 = self.document.add_paragraph()
        para3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        # Format: put letter_id first, then the label (RTL order)
        run3 = para3.add_run(f"{letter_id} : الرقم")
        run3.font.size = Pt(8)
        run3.font.name = 'Tajawal'
        para3.paragraph_format.space_before = Pt(0)
        para3.paragraph_format.space_after = Pt(6)
        para3.paragraph_format.line_spacing = 1.0

        logger.debug(f"Document info header added: {letter_id}")

    def add_besmallah(self, text: str) -> None:
        """
        Add Besmallah (Islamic opening) to the document.

        Args:
            text: The besmallah text (e.g., "بسم الله الرحمن الرحيم")
        """
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = paragraph.add_run(text)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.name = 'Tajawal'

        logger.debug("Besmallah added to document")

    def add_title(self, title: str) -> None:
        """
        Add document title.

        Args:
            title: The document title text
        """
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.style = 'Heading 1'

        run = paragraph.add_run(title)
        run.font.size = Pt(19)
        run.font.bold = True
        # underline
        run.font.underline = True
        run.font.name = 'Tajawal'
        color = RGBColor(0, 0, 0)  # Black color
        run.font.color.rgb = color
        logger.debug(f"Title added: {title}")

    def add_body(self, body: str) -> None:
        """
        Add body content to the document.
        Handles multi-paragraph content by splitting on double newlines.

        Args:
            body: The main content (can contain multiple paragraphs)
        """
        # Split content into paragraphs if multiple newlines exist
        paragraphs = body.split('\n\n') if '\n\n' in body else [body]

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Split single lines by single newlines
            lines = para_text.split('\n')

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                run = paragraph.add_run(line)
                run.font.size = Pt(15)

                run.font.rtl = True 
                run.font.name = 'Tajawal'
                # r = run._r
                # r.get_or_add_rPr().rFonts.set(qn('w:cs'), 'Tajawal') # Sets font for Arabic
                
                # Add spacing between paragraphs
                paragraph.paragraph_format.line_spacing = 1

        logger.debug("Body content added to document")

    def add_footer(self, footer: str) -> None:
        """
        Add footer text to the document.

        Args:
            footer: The footer text
        """
        # Add spacing before footer
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)

        # Add footer content
        para = self.document.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = para.add_run(footer)
        run.font.size = Pt(15)
        run.font.rtl = True
        run.font.name = 'Tajawal'
        run.font.italic = True

        logger.debug("Footer added to document")

    def _download_image_from_drive_url(self, drive_url: str) -> Optional[BytesIO]:
        """
        Download image from Google Drive URL and return as BytesIO.
        Tries two methods: 1) Service account API, 2) Public URL download

        Args:
            drive_url: Google Drive direct link or share link

        Returns:
            BytesIO object containing the image, or None if download fails
        """
        try:
            # Extract file ID from URL
            file_id = None

            if 'drive.google.com' in drive_url:
                if '/file/d/' in drive_url:
                    # Format: https://drive.google.com/file/d/FILE_ID/view
                    file_id = drive_url.split('/file/d/')[1].split('/')[0]
                elif '/document/d/' in drive_url or '/d/' in drive_url:
                    # Format: https://docs.google.com/document/d/FILE_ID/edit or /d/FILE_ID/
                    file_id = drive_url.split('/d/')[1].split('/')[0]
                elif 'id=' in drive_url:
                    # Format: https://drive.google.com/uc?id=FILE_ID
                    file_id = drive_url.split('id=')[1].split('&')[0]

            if not file_id:
                logger.error(f"Could not extract file ID from URL: {drive_url}")
                return None

            logger.debug(f"Extracted file ID: {file_id}")

            # Method 1: Try using service account API (more reliable)
            try:
                from google.oauth2 import service_account
                from googleapiclient.discovery import build
                from googleapiclient.http import MediaIoBaseDownload

                service_account_file = 'automating-letter-creations.json'
                scopes = ['https://www.googleapis.com/auth/drive.readonly']

                creds = service_account.Credentials.from_service_account_file(
                    service_account_file, scopes=scopes
                )
                drive_service = build('drive', 'v3', credentials=creds)

                # Download file using Drive API
                request = drive_service.files().get_media(fileId=file_id)
                image_bytes = BytesIO()
                downloader = MediaIoBaseDownload(image_bytes, request)

                done = False
                while not done:
                    status, done = downloader.next_chunk()

                image_bytes.seek(0)
                logger.info(f"Successfully downloaded image via service account API (size: {len(image_bytes.getvalue())} bytes)")

                # Validate it's an image
                try:
                    img = Image.open(image_bytes)
                    img.verify()
                    image_bytes.seek(0)
                    logger.info(f"Image validated successfully (format: {img.format})")
                    return image_bytes
                except Exception as e:
                    logger.error(f"Downloaded file via API is not a valid image: {e}")
                    return None

            except Exception as api_error:
                logger.warning(f"Service account API download failed: {api_error}. Trying public URL method...")

            # Method 2: Try public URL download
            if 'drive.google.com' in drive_url:
                if file_id:
                    # Create direct download link
                    download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
                else:
                    download_url = drive_url
            else:
                download_url = drive_url

            logger.debug(f"Downloading image from: {download_url}")

            # Download the image
            response = requests.get(download_url, timeout=30)
            response.raise_for_status()

            # Log response details for debugging
            logger.debug(f"Response status: {response.status_code}")
            logger.debug(f"Response content-type: {response.headers.get('content-type', 'unknown')}")
            logger.debug(f"Response content length: {len(response.content)} bytes")

            # Check if response is HTML (common issue with Google Drive)
            if 'text/html' in response.headers.get('content-type', ''):
                logger.error(f"Received HTML instead of image. Google Drive may require authentication or different URL format.")
                logger.debug(f"First 500 chars of response: {response.text[:500]}")
                return None

            # Verify it's an image
            image_bytes = BytesIO(response.content)

            # Try to open with PIL to validate it's an image
            try:
                img = Image.open(image_bytes)
                img.verify()
                logger.info(f"Image downloaded successfully (format: {img.format}, size: {img.size})")
            except Exception as e:
                logger.error(f"Downloaded file is not a valid image: {e}")
                logger.debug(f"First 100 bytes (hex): {response.content[:100].hex()}")
                return None

            # Reset BytesIO position
            image_bytes.seek(0)
            return image_bytes

        except Exception as e:
            logger.error(f"Failed to download image from Drive URL: {e}")
            return None

    def add_signature_section(
        self,
        signature_image_url: str,
        job_title: str,
        name: str,
        section_title: str = "التوقيع"
    ) -> bool:
        """
        Add a signature section to the document with job title, image, and name.
        Layout: Left-aligned, Order: job_title → image → name

        Args:
            signature_image_url: Google Drive URL of the signature image
            job_title: Job title of the signer
            name: Name of the signer
            section_title: Title for the signature section (not used in new layout)

        Returns:
            True if signature section added successfully, False otherwise
        """
        try:
            logger.info(f"Adding signature section for: {name} ({job_title})")

            # Add spacing before signature section
            spacing_para = self.document.add_paragraph()
            spacing_para.paragraph_format.space_before = Pt(18)
            spacing_para.paragraph_format.space_after = Pt(6)

            # Add job title first (left-aligned)
            job_para = self.document.add_paragraph()
            job_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            job_run = job_para.add_run(job_title)
            job_run.font.size = Pt(14)
            job_run.font.bold = True
            job_run.font.name = 'Tajawal'
            job_para.paragraph_format.space_before = Pt(0)
            job_para.paragraph_format.space_after = Pt(6)

            # Download and add signature image
            image_bytes = self._download_image_from_drive_url(signature_image_url)

            if image_bytes:
                # Add image paragraph (left-aligned)
                image_para = self.document.add_paragraph()
                image_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                # Add image with larger size (3 inches width)
                image_run = image_para.add_run()
                image_run.add_picture(image_bytes, width=Inches(3.0))

                image_para.paragraph_format.space_before = Pt(0)
                image_para.paragraph_format.space_after = Pt(6)

                logger.debug("Signature image added successfully")
            else:
                logger.warning("Failed to download signature image, continuing without image")
                # Add a placeholder text if image fails (left-aligned)
                placeholder_para = self.document.add_paragraph()
                placeholder_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                placeholder_run = placeholder_para.add_run("[صورة التوقيع]")
                placeholder_run.font.size = Pt(12)
                placeholder_run.font.italic = True
                placeholder_run.font.name = 'Tajawal'
                placeholder_para.paragraph_format.space_before = Pt(0)
                placeholder_para.paragraph_format.space_after = Pt(6)

            # Add name last (left-aligned)
            name_para = self.document.add_paragraph()
            name_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            name_run = name_para.add_run(name)
            name_run.font.size = Pt(14)
            name_run.font.bold = True
            name_run.font.name = 'Tajawal'
            name_para.paragraph_format.space_before = Pt(0)
            name_para.paragraph_format.space_after = Pt(12)

            logger.info(f"Signature section added successfully for {name}")
            return True

        except Exception as e:
            logger.error(f"Failed to add signature section: {e}")
            return False

    def save(self, output_path: str) -> str:
        """
        Save the document to a file.

        Args:
            output_path: Path where the document will be saved

        Returns:
            The absolute path of the saved file
        """
        if self.document is None:
            raise ValueError("No document loaded. Call create_from_json() first.")

        try:
            # Create directory if it doesn't exist
            os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)

            self.document.save(output_path)
            absolute_path = os.path.abspath(output_path)
            logger.info(f"Document saved successfully to: {absolute_path}")
            return absolute_path
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            raise

    def get_bytes(self) -> BytesIO:
        """
        Get the document as bytes (useful for sending as response/email).

        Returns:
            BytesIO object containing the document
        """
        if self.document is None:
            raise ValueError("No document loaded. Call create_from_json() first.")

        try:
            bytes_io = BytesIO()
            self.document.save(bytes_io)
            bytes_io.seek(0)
            logger.info("Document converted to bytes")
            return bytes_io
        except Exception as e:
            logger.error(f"Failed to convert document to bytes: {e}")
            raise


# # Usage Examples
# if __name__ == "__main__":
#     # Example 1: Create from JSON input (structured data)
#     json_data = {
#         "besmallah": "بسم الله الرحمن الرحيم",
#         "title": "خطاب رسمي",
#         "date": "5 نوفمبر 2025",
#         "body": "محتوى الخطاب يأتي هنا.\n\nهذه فقرة ثانية من الخطاب.",
#         "footer": "مع خالص التحيات والاحترام"
#     }

#     try:
#         # Create document from JSON
#         doc_service = CreateDocument()
#         doc = doc_service.create_from_json(json_data)

#         # Save to file
#         output_path = "output_letter.docx"
#         saved_path = doc_service.save(output_path)
#         print(f"Document saved to: {saved_path}")

#     except Exception as e:
#         print(f"Error (JSON method): {e}")

#     # Example 2: Create from raw letter content (GPT-4 parsing)
#     raw_letter = """بسم الله الرحمن الرحيم
# دعوة لحضور الافتتاح
# صاحب السمو الملكي فايز بن علي الأسمري حفظه الله
# السلام عليكم ورحمة الله وبركاته،
# نود في هذا الخطاب أن نقدم لكم تعريفاً بشركة "نت زيرو"، وهي مشروع اجتماعي
# وطني يرتبط ببرنامج سدرة التابع لوزارة البيئة والمياه والزراعة. تهدف الشركة إلى
# تعزيز الاستدامة البيئية وتحقيق أهداف الحياد الكربوني في المملكة، وذلك من خلال
# تطوير حلول تقنية مبتكرة تدعم التحول نحو مجتمع أكثر وعياً ومسؤولية تجاه البيئة.
# بناءً على ذلك، ندعوكم لحضور الافتتاح، ونأمل أن تشرفونا بحضوركم الكريم
# والمشاركة في هذه المناسبة المهمة، لما تمثله من خطوة فاعلة نحو تحقيق أهداف
# التنمية المستدامة.
# للتنسيق أو الاستفسار يمكنكم التواصل مع مدير العمليات في "نت زيرو"، زياد وائل
# عبد الحميد عبر البريد الإلكتروني: ziadwael@gmail.com أو رقم الجوال:
# 01123808495

# نتقدم إليكم بجزيل الشكر والتقدير على اهتمامكم وتعاونكم، ونتطلع لرؤيتكم
# وتشريفكم.
# وتفضلوا بقبول فائق الاحترام والتقدير،،،"""

#     try:
#         # Create document from raw content (GPT-4 will parse it)
#         doc_service2 = CreateDocument()
#         doc2 = doc_service2.create_from_raw_content(raw_letter)

#         # Save to file
#         output_path2 = "output_letter_parsed.docx"
#         saved_path2 = doc_service2.save(output_path2)
#         print(f"Document saved to: {saved_path2}")

#         # Or get as bytes for sending via email/API response
#         # doc_bytes = doc_service2.get_bytes()

#     except Exception as e:
#         print(f"Error (Raw content method): {e}")


# Create an alias for backward compatibility
EnhancedDOCXService = CreateDocument

# Global service instance
_docx_service = None


def get_enhanced_docx_service() -> CreateDocument:
    """Get the global DOCX document service instance."""
    global _docx_service
    if _docx_service is None:
        _docx_service = CreateDocument()
    return _docx_service
