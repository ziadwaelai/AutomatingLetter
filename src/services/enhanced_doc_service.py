"""
Enhanced DOCX Generation Service using AI and python-docx
Handles DOCX creation from letter content by inserting into template.
"""

import logging
import os
import shutil
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, NamedTuple
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openai
import pytz
from hijri_converter import Hijri, Gregorian

from ..config import get_config
from ..utils import (
    Timer,
    ErrorContext,
    service_error_handler,
    generate_letter_id
)

logger = logging.getLogger(__name__)


class DOCXGenerationResult(NamedTuple):
    """Result of DOCX generation."""
    docx_id: str
    filename: str
    file_path: str
    file_size: int
    generated_at: datetime


class EnhancedDOCXService:
    """Enhanced DOCX generation service using template-based approach."""

    # Styling constants - optimized to fit on one page
    BODY_FONT_NAME = 'Cairo'
    BODY_FONT_SIZE = 11  # Reduced from 17 to fit on one page
    BODY_COLOR = RGBColor(40, 48, 70)  # #283046
    BODY_LINE_SPACING = 1.5  # Reduced from 2.0 to 1.5 for compact layout
    TEXT_ALIGNMENT = WD_ALIGN_PARAGRAPH.JUSTIFY
    PARAGRAPH_INDENT = Inches(0.5)  # Reduced indent from 1.5cm to 0.5 inches

    def __init__(self):
        """Initialize Enhanced DOCX service."""
        self.config = get_config()
        self.output_dir = Path(tempfile.gettempdir()) / "letter_docx"

        # Template path
        self.template_path = Path(__file__).parent.parent.parent / "LetterToPdf" / "Netzero.docx"

        # Create output directory
        self.output_dir.mkdir(exist_ok=True)

        # Initialize OpenAI client
        self.openai_client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

        # Service stats
        self._generation_count = 0
        self._total_size = 0

        logger.info(f"Enhanced DOCX service initialized with template: {self.template_path}")

    def get_current_dates(self) -> Dict[str, str]:
        """Get current date in both Gregorian (KSA timezone) and Hijri formats."""
        try:
            # Get current date in KSA timezone (UTC+3)
            ksa_tz = pytz.timezone('Asia/Riyadh')
            now_ksa = datetime.now(ksa_tz)

            # Format Gregorian date
            gregorian_date = now_ksa.strftime("%Y/%m/%d")

            # Convert to Hijri date
            hijri_date = Gregorian(now_ksa.year, now_ksa.month, now_ksa.day).to_hijri()
            hijri_formatted = f"{hijri_date.year}/{hijri_date.month:02d}/{hijri_date.day:02d}"

            return {
                "gregorian": gregorian_date,
                "hijri": hijri_formatted,
                "gregorian_arabic": f"{now_ksa.day} / {now_ksa.month} / {now_ksa.year}",
                "hijri_arabic": f"{hijri_date.day} / {hijri_date.month} / {hijri_date.year}"
            }
        except Exception as e:
            logger.warning(f"Error getting dates, using fallback: {e}")
            # Fallback to simple date
            now = datetime.now()
            return {
                "gregorian": now.strftime("%Y/%m/%d"),
                "hijri": "N/A",
                "gregorian_arabic": f"{now.day} / {now.month} / {now.year}",
                "hijri_arabic": "N/A"
            }

    def fill_template_with_ai(self, letter_content: str, letter_id: str = "", title: str = "") -> str:
        """
        Fill DOCX template with letter content using AI.

        Args:
            letter_content: Raw letter content
            letter_id: Letter ID for context
            title: Letter title

        Returns:
            Formatted letter body HTML (for structure reference)
        """
        try:
            with Timer("AI template filling"):
                prompt = f"""You are a professional Arabic letter formatter.
Your task is to format letter content for a formal document.

Letter ID: {letter_id}
Title: {title}

Letter Content:
{letter_content}

IMPORTANT FORMATTING INSTRUCTIONS:
1. Split the letter body into separate paragraphs
2. Each logical paragraph or idea should be on its own line
3. Maintain proper Arabic formatting and RTL direction
4. Preserve all original content exactly
5. Return ONLY the formatted paragraphs, one per line
6. DO NOT add any extra content or explanations

Format the letter body content:"""

                response = self.openai_client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {
                            "role": "system",
                            "content": "You are an expert at formatting Arabic formal letters. Format content into properly structured paragraphs."
                        },
                        {
                            "role": "user",
                            "content": prompt
                        }
                    ],
                    temperature=0.3,
                    max_tokens=4000
                )

                formatted_content = response.choices[0].message.content.strip()
                logger.debug(f"AI formatting completed for letter {letter_id}")
                return formatted_content

        except Exception as e:
            logger.error(f"Error in AI template filling: {e}")
            # Return original content if AI fails
            return letter_content


    @service_error_handler
    def generate_docx(
        self,
        title: str,
        content: str,
        recipient: str = "",
        letter_id: str = "",
        document_id: str = ""
    ) -> DOCXGenerationResult:
        """
        Generate DOCX by loading template and inserting letter body content.
        Preserves all header, footer, and styling from the template.

        Args:
            title: Document title
            content: Letter content
            recipient: Letter recipient
            letter_id: Unique letter ID
            document_id: Document ID number

        Returns:
            DOCXGenerationResult with file information
        """
        try:
            with Timer("DOCX generation"):
                # Check if template exists
                if not self.template_path.exists():
                    raise FileNotFoundError(f"Template not found: {self.template_path}")

                # Load template document
                logger.info(f"Loading template from: {self.template_path}")
                doc = Document(str(self.template_path))

                # Set document-level direction to RTL
                sectPr = doc.sections[0]._sectPr
                # Add bidi element for RTL (Right-to-Left)
                existing_bidi = sectPr.find(qn('w:bidi'))
                if existing_bidi is None:
                    bidi = OxmlElement('w:bidi')
                    sectPr.insert(0, bidi)

                # Format content using AI to split into proper paragraphs
                formatted_content = self.fill_template_with_ai(content, letter_id, title)

                # Split into paragraphs
                paragraphs_text = formatted_content.split('\n')

                # Clear existing body paragraphs (keep only header/footer)
                # Get the main body and insert after template content
                # For template, we insert at the end of document

                # Add spacing after template content
                doc.add_paragraph()

                # Insert formatted letter body paragraphs
                for para_text in paragraphs_text:
                    para_text = para_text.strip()
                    if para_text:  # Skip empty lines
                        para = doc.add_paragraph(para_text)
                        para.alignment = self.TEXT_ALIGNMENT
                        para.paragraph_format.right_to_left = True  # RTL (Right-to-Left)
                        para.paragraph_format.first_line_indent = self.PARAGRAPH_INDENT
                        para.paragraph_format.line_spacing = self.BODY_LINE_SPACING
                        para.paragraph_format.space_after = Pt(6)  # Reduced from 12 to 6
                        para.paragraph_format.space_before = Pt(0)  # No space before

                        # Set RTL at XML level
                        pPr = para._element.get_or_add_pPr()
                        # Add bidi element for RTL
                        existing_bidi = pPr.find(qn('w:bidi'))
                        if existing_bidi is None:
                            bidi = OxmlElement('w:bidi')
                            pPr.insert(0, bidi)

                        # Apply font styling to all runs
                        for run in para.runs:
                            run.font.size = Pt(self.BODY_FONT_SIZE)
                            run.font.color.rgb = self.BODY_COLOR
                            run.font.name = self.BODY_FONT_NAME

                # Generate filename and save
                docx_id = str(uuid.uuid4())
                filename = f"letter_{letter_id}_{docx_id}.docx"
                file_path = self.output_dir / filename

                doc.save(str(file_path))

                file_size = file_path.stat().st_size

                # Update stats
                self._generation_count += 1
                self._total_size += file_size

                result = DOCXGenerationResult(
                    docx_id=docx_id,
                    filename=filename,
                    file_path=str(file_path),
                    file_size=file_size,
                    generated_at=datetime.now()
                )

                logger.info(f"DOCX generated successfully from template: {filename} ({file_size} bytes)")
                return result

        except Exception as e:
            logger.error(f"Error generating DOCX: {e}", exc_info=True)
            raise

    def get_service_stats(self) -> Dict[str, Any]:
        """Get service statistics."""
        return {
            "generation_count": self._generation_count,
            "total_size_bytes": self._total_size,
            "output_directory": str(self.output_dir)
        }


# Singleton instance
_docx_service_instance: Optional[EnhancedDOCXService] = None


def get_enhanced_docx_service() -> EnhancedDOCXService:
    """Get or create Enhanced DOCX service instance."""
    global _docx_service_instance
    if _docx_service_instance is None:
        _docx_service_instance = EnhancedDOCXService()
    return _docx_service_instance
